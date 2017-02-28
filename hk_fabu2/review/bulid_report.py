# -*- coding:utf-8 -*-
import os,django,sys
BASE_DIR =os.path.dirname(os.path.abspath(__file__))
sys.path.append(BASE_DIR)

# print sys.path

# from models import ReviewAction
from docx import Document
from docx.shared import Inches
import time
import xlwt

# curr_date_str = time.strftime('%Y%m%d%H%M%S')

itme_list =[u'软件系统简称及称:',u'版本发布申请人:',u'发布类型:',u'发布版本号:',\
u'发布模块名:',u'发布文件名:',u'发布时间:',u'发布状态:',u'发布主机:']

item_values=['tomcat_omweb',u'张三',u'全新发布','3.0.1','tomcat_omweb','tomcat_omweb.zip',\
'2016-12-05 12:00:00',u'审核通过,发布完成',u'10.0.0.1,10.0.0.2']

detail_list=[u'序号',u'类型',u'详细描述',u'开发作者',u'所属子系统',u'发布者',u'备注',u'测试工程师']
detail_values=[[u'1',u'bug修复',u'修改注册码',u'张三',u'user',u'张三',u'无',u'李四'],\
[u'2',u'需求变更',u'去掉邀请码字段',u'王五',u'microapi',u'王五',u'无',u'赵六'],\
[u'3',u'新增功能',u'增加验证',u'李七',u'user',u'李七',u'无',u'罗八']]

check_list = [u'审批人',u'审批意见',u'审批时间']
check_values = [u'研发主管张三',u'同意发布',u'2016-12-05 16:03:12']

def bulid_docs(publish_id,base_itemlist,base_valuelist,review_itemlist,review_valuelist,test_itemlist,test_valuelist,publish_detail,test_detail):
	document = Document()
	document.add_heading(u'版本发布报告', 0)
	itme_list =[u'软件系统简称及称:',u'版本发布申请人:',u'发布类型:',u'发布版本号:',\
	u'发布模块名:',u'发布文件名:',u'发布时间:',u'发布状态:',u'发布主机:']

	# item_obj = models.ReviewAction.objects.get(publish_id=publish_id)
	# item_values =[u'华康web3.0系统']
	# item_values.append(item_obj.publish_user,\
	# 	u'全新发布',\
	# 	u'3.0.1',\
	# 	item_obj.publish_module,\
	# 	item_obj.publish_filename,\
	# 	item_obj.create_time,\
	# 	item_obj.publish_status,\
	# 	item_obj.publish_serverlist)
	item_values=['tomcat_omweb',u'张三',u'全新发布','3.0.1','tomcat_omweb','tomcat_omweb.zip',\
	'2016-12-05 12:00:00',u'审核通过,发布完成',u'10.0.0.1,10.0.0.2']

	detail_list=[u'序号',u'类型',u'详细描述',u'开发作者',u'所属子系统',u'发布者',u'备注',u'测试工程师']
	detail_values=[[u'1',u'bug修复',u'修改注册码',u'张三',u'user',u'张三',u'无',u'李四'],\
	[u'2',u'需求变更',u'去掉邀请码字段',u'王五',u'microapi',u'王五',u'无',u'赵六'],\
	[u'3',u'新增功能',u'增加验证',u'李七',u'user',u'李七',u'无',u'罗八']]

	check_list = [u'审批人',u'审批意见',u'审批时间']
	check_values = [u'研发主管张三',u'同意发布',u'2016-12-05 16:03:12']

	# p = document.add_paragraph('A plain paragraph having some ')
	# p.add_run('bold').bold = True
	# p.add_run(' and some ')
	# p.add_run('italic.').italic = True

	document.add_heading(u'基本信息', level=3)
	# document.add_paragraph('Intense quote', style='IntenseQuote')

	# document.add_paragraph(
	#     'first item in unordered list', style='ListBullet'
	# )
	# document.add_paragraph(
	#     'first item in ordered list', style='ListNumber'
	# )

	# document.add_picture('monty-truth.png', width=Inches(1.25))


	# document.add_paragraph(u'基本信息', style='ListNumber')

	base_info_tables = document.add_table(rows=1, cols=2)
	for i in range(9):
		base_row_cell = base_info_tables.add_row().cells
		base_row_cell[0].text = itme_list[i]
		base_row_cell[1].text = item_values[i]

	#document.add_section(start_type=2)手动分页
	# document.add_paragraph(u'基本信息',style='ListNumber')
	document.add_heading(u'详细信息说明', level=3)

	detail_info_tables = document.add_table(rows=1, cols=8)
	detail_table_title = detail_info_tables.rows[0].cells
	for i in range(len(detail_list)):
		detail_table_title[i].text = detail_list[i]

	# detail_table_title[0].text = 
	for i in range(3):
		detail_row_cell = detail_info_tables.add_row().cells
		detail_row_cell[0].text = detail_values[i][0]
		detail_row_cell[1].text = detail_values[i][1]
		detail_row_cell[2].text = detail_values[i][2]
		detail_row_cell[3].text = detail_values[i][3]
		detail_row_cell[4].text = detail_values[i][4]
		detail_row_cell[5].text = detail_values[i][5]
		detail_row_cell[6].text = detail_values[i][6]
		detail_row_cell[7].text = detail_values[i][7]


	document.add_heading(u'审批意见', level=3)

	review_table = document.add_table(rows=1,cols=2)
	review_title = review_table.rows[0].cells[0].text = u'审批意见'
	for i in range(3):
		review_row_cell = review_table.add_row().cells
		review_row_cell[0].text = check_list[i]
		review_row_cell[1].text = check_values[i]



	document.add_heading(u'测试信息说明', level=3)


	# table = document.add_table(rows=1, cols=3)
	# hdr_cells = table.rows[0].cells
	# hdr_cells[0].text = U'序号'
	# hdr_cells[1].text = u'需求编号'
	# hdr_cells[2].text = u'功能名称'
	# for i in range(2):
	# 	row_cells = table.add_row().cells
	# 	row_cells[0].text = '1'
	# 	row_cells[1].text = '20161205'
	# 	row_cells[2].text = u'api接口'


	# for item in recordset:
	#     row_cells = table.add_row().cells
	#     row_cells[0].text = str(item.qty)
	#     row_cells[1].text = str(item.id)
	#     row_cells[2].text = item.desc

	document.add_page_break()
	curr_date_str = time.strftime('%Y%m%d%H%M%S')
	document.save('/home/fabu/hk_fabu2/review/download/doc/'+curr_date_str+'.doc')
	
	return 'doc'+curr_date_str

def bulid_excel(publish_id,base_itemlist,base_valuelist,review_itemlist,review_valuelist,test_itemlist,test_valuelist,publish_detail,test_detail):
    w = xlwt.Workbook(encoding='utf-8')
    curr_date_str = time.strftime('%Y%m%d%H%M%S')

    ws = w.add_sheet(u'版本发布报告')
    #设置列宽
    fir_col = ws.col(0)
    sec_col = ws.col(1)

    fir_col.width =256*30
    sec_col.width =256*90

    #第一部分,表格标题
    # ws.write(0,0,u'版本发布报告',set_style('Times New Roman',220,True))
    ws.write_merge(0,0,0,1,u'版本发布报告',set_style(u'微软雅黑',400,True))
    #第二部分，基本信息
    ws.write_merge(1,1,0,1,u'版本基本信息',set_style('Times New Roman',320,True))
    for i in range(len(base_itemlist)):
        ws.write(i+2,0,base_itemlist[i],set_style('Times New Roman',320,True))
        ws.write(i+2,1,base_valuelist[i],set_style('Times New Roman',320,False))

    # for j in range(len(value_list)):
    #     ws.write(j+1,1,value_list[j])
    #
    #第三部分，审核信息
    ws.write_merge(len(base_itemlist)+2,len(base_itemlist)+2,0,1,u'审核发布信息',set_style('Times New Roman',320,True))
    for i in range(len(review_itemlist)):
        ws.write(i+2+len(base_itemlist)+1,0,review_itemlist[i],set_style('Times New Roman',320,True))
        ws.write(i+2+len(base_itemlist)+1,1,review_valuelist[i],set_style('Times New Roman',320,False))

    check_ws = w.add_sheet(u'审核发布详情')
    check_ws.write_merge(0,0,0,1,u'审核发布详情',set_style(u'微软雅黑',400,True))






    #第四部分，测试信息
    ws.write_merge(len(base_itemlist)+len(review_itemlist)+3,len(base_itemlist)+len(review_itemlist)+3,0,1,u'测试信息',set_style('Times New Roman',320,True))
    for i in range(len(test_itemlist)):
        ws.write(i+2+len(base_itemlist)+len(review_itemlist)+1+1,0,test_itemlist[i],set_style('Times New Roman',320,True))
        ws.write(i+2+len(base_itemlist)+len(review_itemlist)+1+1,1,test_valuelist[i],set_style('Times New Roman',320,False)) 
    
    test_ws = w.add_sheet(u'测试发布详情')
    test_ws.write_merge(0,0,0,1,u'测试发布详情',set_style(u'微软雅黑',400,True)) 




    w.save('/home/fabu/hk_fabu2/review/download/xls/'+curr_date_str+'.xls')

    # import pandas
    # df=pandas.read_html(str(publish_detail).replace('\n',''))
    # bb = pandas.ExcelWriter('/home/fabu/hk_fabu2/review/download/xls/'+curr_date_str+'.xls')
    # df[0].to_excel(bb,u'审核发布详情')
    # bb.save()
    return 'xls'+curr_date_str

def set_style(name,height,bold=False):
    style = xlwt.XFStyle()
    font = xlwt.Font() # 为样式创建字体
    font.name = name # 'Times New Roman'
    font.bold = bold
    font.color_index = 4
    font.height = height
    #height中，220-720分别对应字体号11-36

    borders= xlwt.Borders()
    borders.left= 1
    borders.right= 1
    borders.top= 1
    borders.bottom= 1

    alignment = xlwt.Alignment()
    alignment.horz = xlwt.Alignment.HORZ_CENTER # 垂直对齐
    alignment.vert = xlwt.Alignment.VERT_CENTER # 水平对齐
    alignment.wrap = xlwt.Alignment.WRAP_AT_RIGHT # 自动换行

    style.font = font
    style.borders = borders
    style.alignment = alignment

    return style


# bulid_docs(1)

